import streamlit as st
import asyncio
import pandas as pd
from typing import List
import os
import json
from datetime import datetime
from dotenv import load_dotenv
from multi_agent_system import MultiAgentSystem, SearchResult
from paper_writer import ResearchPaperWriter
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# Load environment variables
load_dotenv()

# Configure page settings
st.set_page_config(
    page_title="Research Assistant",
    page_icon="🔍",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for improved UI
st.markdown("""
<style>
    /* Main app styling */
    .main {
        background-color: #0E1117;
    }
    
    /* Headers */
    h1 {
        color: #FFFFFF !important;
        font-family: 'Segoe UI', sans-serif;
        font-weight: 700;
        letter-spacing: -0.5px;
        padding-bottom: 10px;
    }
    
    h2 {
        color: #E0E0E0 !important;
        font-family: 'Segoe UI', sans-serif;
        font-weight: 600;
        padding-top: 20px;
        padding-bottom: 10px;
    }
    
    h3 {
        color: #B0B0B0 !important;
        font-family: 'Segoe UI', sans-serif;
        font-weight: 500;
    }
    
    /* Paragraphs and text */
    p, li {
        font-size: 15px;
        line-height: 1.7;
        color: #D0D0D0;
    }
    
    /* Buttons */
    .stButton>button {
        border-radius: 8px;
        font-weight: 600;
        font-size: 16px;
        padding: 12px 24px;
        transition: all 0.3s ease;
        border: none;
    }
    
    .stButton>button:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(103, 126, 234, 0.4);
    }
    
    /* Input fields */
    .stTextInput>div>div>input {
        border-radius: 8px;
        border: 2px solid #667eea;
        padding: 12px;
        font-size: 15px;
        background-color: #1E1E1E;
        color: #FFFFFF;
    }
    
    .stTextInput>div>div>input:focus {
        border-color: #764ba2;
        box-shadow: 0 0 0 2px rgba(118, 75, 162, 0.2);
    }
    
    /* Metrics */
    [data-testid="stMetricValue"] {
        font-size: 28px;
        font-weight: 700;
        color: #667eea;
    }
    
    [data-testid="stMetricLabel"] {
        font-size: 14px;
        font-weight: 500;
        color: #B0B0B0;
    }
    
    /* Tabs */
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
        background-color: #1A1A1A;
        padding: 10px;
        border-radius: 10px;
    }
    
    .stTabs [data-baseweb="tab"] {
        border-radius: 8px;
        padding: 12px 20px;
        font-weight: 600;
        font-size: 15px;
        background-color: #2A2A2A;
        color: #B0B0B0;
    }
    
    .stTabs [aria-selected="true"] {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
    }
    
    /* Expanders */
    .streamlit-expanderHeader {
        background-color: #1E1E1E;
        border-radius: 8px;
        font-weight: 600;
        color: #E0E0E0;
        border-left: 4px solid #667eea;
    }
    
    .streamlit-expanderHeader:hover {
        background-color: #2A2A2A;
    }
    
    /* Dataframes */
    .dataframe {
        border-radius: 8px;
        overflow: hidden;
    }
    
    /* Progress bar */
    .stProgress > div > div > div > div {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
    }
    
    /* Success/Info/Warning/Error boxes */
    .stSuccess, .stInfo, .stWarning, .stError {
        border-radius: 8px;
        padding: 15px;
        border-left: 4px solid;
    }
    
    /* Sidebar */
    section[data-testid="stSidebar"] {
        background-color: #0A0A0A;
    }
    
    /* Download buttons */
    .stDownloadButton>button {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border: none;
        border-radius: 8px;
        font-weight: 600;
    }
    
    /* Dividers */
    hr {
        border-color: #2A2A2A;
        margin: 30px 0;
    }
</style>
""", unsafe_allow_html=True)

def init_session_state():
    """Initialize session state variables"""
    if 'search_history' not in st.session_state:
        st.session_state.search_history = []
    if 'analysis_progress' not in st.session_state:
        st.session_state.analysis_progress = 0

def get_score_badge(score, label):
    """Generate color-coded badge for scores"""
    try:
        score_val = float(score)
        if score_val >= 8:
            color = "#00C853"  # Green
            emoji = "🟢"
        elif score_val >= 6:
            color = "#FFA726"  # Orange
            emoji = "🟡"
        else:
            color = "#EF5350"  # Red
            emoji = "🔴"
        
        return f"""
        <div style="display: inline-block; background-color: {color}; color: white; padding: 5px 12px; 
                    border-radius: 15px; margin: 3px; font-weight: bold; font-size: 14px;">
            {emoji} {label}: {score_val:.1f}/10
        </div>
        """
    except:
        return f"<span>{label}: {score}</span>"

def parse_literature_review(review_text):
    """Parse literature review into collapsible sections"""
    sections = {}
    current_section = "Introduction"
    current_content = []
    
    lines = review_text.split('\n')
    for line in lines:
        # Detect section headers
        if line.startswith('##'):
            # Save previous section
            if current_content:
                sections[current_section] = '\n'.join(current_content)
            # Start new section
            current_section = line.replace('#', '').strip()
            current_content = []
        else:
            current_content.append(line)
    
    # Save last section
    if current_content:
        sections[current_section] = '\n'.join(current_content)
    
    return sections if sections else {"Full Review": review_text}

def convert_markdown_to_docx(markdown_text: str, title: str = "Research Paper") -> bytes:
    """Convert markdown text to Word document (.docx) format"""
    try:
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Split markdown into lines
        lines = markdown_text.split('\n')
        
        for line in lines:
            line = line.strip()
            
            if not line:
                continue
            
            # Handle headers
            if line.startswith('# '):
                # Title (H1)
                p = doc.add_heading(line[2:], level=1)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif line.startswith('## '):
                # Section (H2)
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                # Subsection (H3)
                doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                # Sub-subsection (H4)
                doc.add_heading(line[5:], level=4)
            elif line.startswith('- ') or line.startswith('* '):
                # Bullet point
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('> '):
                # Block quote
                p = doc.add_paragraph(line[2:])
                p.style = 'Intense Quote'
            else:
                # Regular paragraph
                doc.add_paragraph(line)
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()
        
    except Exception as e:
        st.error(f"Error converting to Word: {str(e)}")
        return None

async def run_async_operations(system, search_query, max_results=3):
    """Handle async complete analysis operation"""
    # Always run full pipeline: Paper Collection → Literature Review → Gap Analysis → Hypothesis Generation
    complete_results = await system.run_complete_analysis(search_query)
    return {"type": "complete", "data": complete_results}

def main():
    """Main application function"""
    st.markdown("""
    <div style="text-align: center; padding: 30px 20px; margin-bottom: 20px; 
                background: linear-gradient(135deg, #0f2027 0%, #203a43 50%, #2c5364 100%);
                border-radius: 12px; border: 1px solid rgba(255,255,255,0.1);">
        <h1 style="font-size: 42px; margin-bottom: 8px; color: #ffffff; font-weight: 700; letter-spacing: 0.5px;">
            ResearchForge
        </h1>
        <p style="font-size: 16px; color: #a0aec0; font-weight: 500; margin-bottom: 15px;">
            Autonomous AI Research Assistant
        </p>
        <p style="font-size: 14px; color: #cbd5e0; font-weight: 400; line-height: 1.8;">
            Automated Literature Review • Gap Analysis • Hypothesis Generation<br/>
            Experiment Design • Dataset Recommendations • AI Research Paper Writer
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    init_session_state()

    if 'system' not in st.session_state:
        with st.spinner("🔄 Initializing Research Assistant... Please wait..."):
            st.session_state.system = MultiAgentSystem()
            st.success("✅ System initialized! Models will load when needed.")
    
    system = st.session_state.system

    # Sidebar configuration
    with st.sidebar:
        st.markdown("""
        <div style="background: linear-gradient(135deg, #2c5364 0%, #203a43 100%); 
                    padding: 18px; border-radius: 8px; margin-bottom: 20px; border: 1px solid rgba(255,255,255,0.1);">
            <h2 style="color: white; margin: 0; text-align: center; font-size: 20px; font-weight: 600;">Configuration</h2>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("### Analysis Pipeline")
        st.markdown("""
        <div style="background-color: #1e293b; padding: 15px; border-radius: 8px; border-left: 3px solid #3b82f6;">
            <p style="color: #e2e8f0; font-size: 14px; margin-bottom: 10px; font-weight: 600;">What happens when you search:</p>
            <p style="color: #cbd5e0; margin: 6px 0; font-size: 13px;">1. 📄 Collect 10-15 research papers</p>
            <p style="color: #cbd5e0; margin: 6px 0; font-size: 13px;">2. 📚 Literature Review (GPT-3.5)</p>
            <p style="color: #cbd5e0; margin: 6px 0; font-size: 13px;">3. 🔍 Gap Analysis (GPT-4)</p>
            <p style="color: #cbd5e0; margin: 6px 0; font-size: 13px;">4. 💡 Hypothesis Generation (GPT-4)</p>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("")

        # API Status
        st.markdown("### API Status")
        if os.getenv('OPENAI_API_KEY') and os.getenv('GOOGLE_SEARCH_ENGINE_ID'):
            st.markdown("""
            <div style="background-color: #064e3b; padding: 12px; border-radius: 8px; border-left: 3px solid #10b981;">
                <p style="color: #d1fae5; margin: 0; font-weight: 600; font-size: 14px;">✅ All APIs configured</p>
            </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown("""
            <div style="background-color: #7f1d1d; padding: 12px; border-radius: 8px; border-left: 3px solid #ef4444;">
                <p style="color: #fecaca; margin: 0; font-weight: 600; font-size: 14px;">❌ APIs not configured. Please check .env file</p>
            </div>
            """, unsafe_allow_html=True)
        
        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown("---")
        
        st.markdown("""
        <div style="background-color: #1e293b; padding: 15px; border-radius: 8px;">
            <p style="color: #60a5fa; font-weight: 600; margin-bottom: 10px; font-size: 14px;">Powered by:</p>
            <p style="color: #e2e8f0; margin: 5px 0; padding-left: 10px; font-size: 13px;">• OpenAI GPT-4 & GPT-3.5</p>
            <p style="color: #e2e8f0; margin: 5px 0; padding-left: 10px; font-size: 13px;">• Sentence-BERT</p>
        </div>
        """, unsafe_allow_html=True)

        # Add info about model loading
        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown("""
        <div style="background-color: #134e4a; padding: 12px; border-radius: 8px; border-left: 3px solid #14b8a6;">
            <p style="color: #5eead4; margin: 0; font-size: 13px;">💡 <b>Note:</b> Heavy AI models load on first use. Subsequent uses are instant!</p>
        </div>
        """, unsafe_allow_html=True)


    # Main search interface
    st.markdown("")
    st.markdown("")
    
    # Show current analysis status
    if st.session_state.get('complete_data'):
        current_topic = st.session_state.get('search_query', 'Unknown')
        st.markdown(f"""
        <div style="background-color: rgba(16, 185, 129, 0.1); padding: 12px 20px; border-radius: 6px; 
                    border-left: 3px solid #10b981; margin-bottom: 20px;">
            <p style="color: #6ee7b7; margin: 0; font-size: 14px; font-weight: 500;">
                ✓ Analysis completed for: <span style="color: #d1fae5;">{current_topic}</span>
            </p>
        </div>
        """, unsafe_allow_html=True)
        
        col1, col2 = st.columns([3, 1])
        with col2:
            if st.button("🔄 New Analysis", type="secondary", use_container_width=True):
                # Clear all session state
                for key in ['complete_data', 'search_query', 'selected_hypothesis', 'selected_gap', 
                           'experiment_results', 'selected_experiment', 'generated_paper']:
                    if key in st.session_state:
                        del st.session_state[key]
                st.rerun()
    
    # Clean, minimal search input
    st.markdown("""
    <p style="color: #cbd5e0; font-size: 15px; font-weight: 500; margin-bottom: 10px;">
        🔍 Enter Your Research Topic
    </p>
    """, unsafe_allow_html=True)
    
    search_query = st.text_input(
        "Research Topic",
        placeholder="e.g., machine learning for drug discovery, quantum computing applications...",
        label_visibility="collapsed",
        key="search_input"
    )
    
    search_button = st.button("🚀 Start Analysis", type="primary", use_container_width=True)
    
    st.markdown("<br>", unsafe_allow_html=True)

    # Handle search
    if search_button and search_query:
        # Create progress bar
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        try:
            # Phase 1: Paper Collection
            status_text.markdown("""
            <div style="padding: 10px 15px; border-left: 3px solid #3b82f6; background-color: rgba(59, 130, 246, 0.1);">
                <p style="color: #93c5fd; margin: 0; font-size: 14px;">📄 Phase 1/4: Collecting research papers...</p>
            </div>
            """, unsafe_allow_html=True)
            progress_bar.progress(10)
            
            system = MultiAgentSystem()
            
            # Phase 2: Literature Review
            status_text.markdown("""
            <div style="padding: 10px 15px; border-left: 3px solid #10b981; background-color: rgba(16, 185, 129, 0.1);">
                <p style="color: #6ee7b7; margin: 0; font-size: 14px;">📚 Phase 2/4: Analyzing literature...</p>
            </div>
            """, unsafe_allow_html=True)
            progress_bar.progress(30)
            
            # Run analysis (this includes all phases)
            result = asyncio.run(run_async_operations(
                system, 
                search_query,
                max_results=5
            ))
            
            # Phase 3: Gap Analysis
            status_text.markdown("""
            <div style="padding: 10px 15px; border-left: 3px solid #f59e0b; background-color: rgba(245, 158, 11, 0.1);">
                <p style="color: #fcd34d; margin: 0; font-size: 14px;">🔍 Phase 3/4: Identifying research gaps...</p>
            </div>
            """, unsafe_allow_html=True)
            progress_bar.progress(65)
            
            # Phase 4: Hypothesis Generation
            status_text.markdown("""
            <div style="padding: 10px 15px; border-left: 3px solid #8b5cf6; background-color: rgba(139, 92, 246, 0.1);">
                <p style="color: #c4b5fd; margin: 0; font-size: 14px;">💡 Phase 4/4: Generating hypotheses...</p>
            </div>
            """, unsafe_allow_html=True)
            progress_bar.progress(90)
            
            # Complete
            complete_data = result["data"]
            st.session_state['complete_data'] = complete_data
            st.session_state['search_query'] = search_query
            
            progress_bar.progress(100)
            status_text.markdown("""
            <div style="padding: 12px 20px; border-left: 3px solid #10b981; background-color: rgba(16, 185, 129, 0.15);">
                <p style="color: #6ee7b7; margin: 0; font-size: 15px; font-weight: 600;">
                    ✓ Analysis Complete
                </p>
            </div>
            """, unsafe_allow_html=True)
            
            st.balloons()
                
        except Exception as e:
            st.error(f"An error occurred: {str(e)}")
            st.error(f"Error details: {type(e).__name__}")
    
    # Display tabs if we have complete data in session state
    if st.session_state.get('complete_data'):
        complete_data = st.session_state['complete_data']
        search_query = st.session_state.get('search_query', '')
        
        # Create tabs for different sections
        tab1, tab2, tab3, tab4, tab5 = st.tabs([
            "📚 Literature Review", 
            "🔍 Gap Analysis", 
            "💡 Hypotheses",
            "🧪 Experiment Design",
            "📝 AI Paper Writer"
        ])
        
        with tab1:
            st.header("Literature Review")
            
            # Show papers collected
            papers = complete_data.get("papers", [])
            if papers:
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("📄 Papers Found", len(papers))
                with col2:
                    arxiv_count = sum(1 for p in papers if hasattr(p, 'arxiv_url') and p.arxiv_url)
                    st.metric("📚 arXiv Papers", arxiv_count)
                with col3:
                    google_count = len(papers) - arxiv_count
                    st.metric("🔍 Google Scholar", google_count)
                
                st.markdown("---")
            
            # Show literature review with collapsible sections
            if complete_data.get("literature_review"):
                review_text = complete_data["literature_review"]
                
                # Download button for literature review
                col1, col2 = st.columns([3, 1])
                with col2:
                    st.download_button(
                        label="📥 Download as TXT",
                        data=review_text,
                        file_name=f"literature_review_{datetime.now().strftime('%Y%m%d')}.txt",
                        mime="text/plain",
                        use_container_width=True
                    )
                
                # Parse into sections
                sections = parse_literature_review(review_text)
                
                if len(sections) > 1:
                    st.markdown("### 📖 Review Sections (Click to Expand)")
                    for section_name, section_content in sections.items():
                        with st.expander(f"📑 {section_name}", expanded=(section_name == "Introduction" or section_name == "Full Review")):
                            st.markdown(section_content)
                else:
                    st.markdown(review_text)
            else:
                st.warning("No literature review available")
            
            # Display collected papers with links
            if papers:
                st.markdown("---")
                st.subheader("📚 Research Papers Analyzed")
                st.info(f"💡 {len(papers)} papers were collected and analyzed for this review")
                
                for i, paper in enumerate(papers, 1):
                    with st.expander(f"📄 Paper {i}: {paper.title}", expanded=False):
                        if hasattr(paper, 'authors') and paper.authors:
                            st.markdown(f"**✍️ Authors:** {', '.join(paper.authors[:3])}{'...' if len(paper.authors) > 3 else ''}")
                        
                        if hasattr(paper, 'published') and paper.published:
                            st.markdown(f"**📅 Published:** {paper.published}")
                        
                        if hasattr(paper, 'abstract') and paper.abstract:
                            st.markdown(f"**📝 Abstract:**")
                            st.markdown(paper.abstract[:500] + "..." if len(paper.abstract) > 500 else paper.abstract)
                        elif hasattr(paper, 'snippet') and paper.snippet:
                            st.markdown(f"**📝 Summary:**")
                            st.markdown(paper.snippet)
                        
                        # Links
                        links = []
                        if hasattr(paper, 'arxiv_url') and paper.arxiv_url:
                            links.append(f"[📄 arXiv]({paper.arxiv_url})")
                        if hasattr(paper, 'pdf_url') and paper.pdf_url:
                            links.append(f"[📥 PDF]({paper.pdf_url})")
                        if hasattr(paper, 'link') and paper.link:
                            links.append(f"[🔗 Source]({paper.link})")
                        
                        if links:
                            st.markdown("**🔗 Links:** " + " | ".join(links))
            else:
                st.warning("No papers were collected for this analysis")
        
        with tab2:
            st.header("Research Gap Analysis")
            gap_data = complete_data.get("gap_analysis", {})
            
            # Show which model was used
            gaps_dict = gap_data.get("gaps", {})
            if gaps_dict.get("source"):
                if "Combined" in gaps_dict["source"]:
                    st.success(f"🤖 Analysis by: {gaps_dict['source']}")
                elif "GPT-4" in gaps_dict["source"]:
                    st.info(f"🤖 Analysis by: {gaps_dict['source']}")
                else:
                    st.success(f"🤖 Analysis by: {gaps_dict['source']}")
            
            # Show summary metrics
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("Papers Analyzed", gap_data.get("papers_analyzed", 0))
            with col2:
                total_gaps = sum(len(gaps_dict.get(k, [])) for k in ["knowledge_gaps", "methodological_gaps", "dataset_gaps", "temporal_gaps"])
                st.metric("Total Gaps Found", total_gaps)
            with col3:
                st.metric("Categories", 4)
            with col4:
                # Download button
                gap_text = gap_data.get("formatted_output", "No gap analysis available")
                st.download_button(
                    label="📥 Download",
                    data=gap_text,
                    file_name=f"gap_analysis_{datetime.now().strftime('%Y%m%d')}.txt",
                    mime="text/plain",
                    use_container_width=True
                )
            
            # Display formatted gap analysis
            st.markdown(gap_data.get("formatted_output", "No gap analysis available"))
        
        with tab3:
            st.header("Research Hypotheses")
            hyp_data = complete_data.get("hypotheses", {})
            
            if hyp_data.get("hypotheses"):
                hypotheses_list = hyp_data.get("hypotheses", [])
                
                # Show summary metrics
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("Hypotheses Generated", len(hypotheses_list))
                with col2:
                    avg_score = sum(h.get("overall_score", 0) for h in hypotheses_list) / max(len(hypotheses_list), 1)
                    st.metric("Avg Overall Score", f"{avg_score:.2f}")
                with col3:
                    avg_novelty = sum(h.get("novelty_score", 0) for h in hypotheses_list) / max(len(hypotheses_list), 1)
                    st.metric("Avg Novelty Score", f"{avg_novelty:.2f}")
                
                st.markdown("---")
                
                # Side-by-side comparison table
                st.markdown("### 📊 Hypothesis Comparison Table")
                
                comparison_data = []
                for idx, hyp in enumerate(hypotheses_list, 1):
                    comparison_data.append({
                        "#": idx,
                        "Title": hyp.get('title', 'Untitled')[:50] + '...' if len(hyp.get('title', '')) > 50 else hyp.get('title', 'Untitled'),
                        "Overall": f"{hyp.get('overall_score', 0):.1f}",
                        "Novelty": f"{hyp.get('novelty_score', 0):.1f}",
                        "Impact": f"{hyp.get('impact_score', 0):.1f}",
                        "Feasibility": f"{hyp.get('feasibility_score', 0):.1f}",
                        "Gap": hyp.get('gap_addressed', 'N/A')[:40] + '...' if len(hyp.get('gap_addressed', '')) > 40 else hyp.get('gap_addressed', 'N/A')
                    })
                
                comparison_df = pd.DataFrame(comparison_data)
                st.dataframe(comparison_df, use_container_width=True, hide_index=True)
                
                # Download CSV button
                csv = comparison_df.to_csv(index=False)
                st.download_button(
                    label="📥 Download Comparison as CSV",
                    data=csv,
                    file_name=f"hypothesis_comparison_{datetime.now().strftime('%Y%m%d')}.csv",
                    mime="text/csv",
                    use_container_width=True
                )
                
                st.markdown("---")
                
                # Display hypotheses with selection and color-coded badges
                st.markdown("### 💡 Detailed Hypotheses (Click to Expand)")
                
                for idx, hyp in enumerate(hypotheses_list, 1):
                    with st.expander(f"{'🔴' if hyp.get('priority') == 'HIGH' else '🟡'} Hypothesis {idx}: {hyp.get('title', 'Untitled')}", expanded=(idx==1)):
                        # Color-coded score badges
                        st.markdown("**📊 Performance Scores:**", unsafe_allow_html=True)
                        
                        badges_html = ""
                        badges_html += get_score_badge(hyp.get('overall_score', 0), 'Overall')
                        badges_html += get_score_badge(hyp.get('novelty_score', 0), 'Novelty')
                        badges_html += get_score_badge(hyp.get('impact_score', 0), 'Impact')
                        badges_html += get_score_badge(hyp.get('feasibility_score', 0), 'Feasibility')
                        
                        st.markdown(badges_html, unsafe_allow_html=True)
                        st.markdown("")
                        
                        if hyp.get('gap_addressed'):
                            st.markdown(f"**🎯 Gap Addressed:** {hyp['gap_addressed']}")
                        
                        if hyp.get('problem_statement'):
                            st.markdown(f"**❓ Problem Statement:** {hyp['problem_statement']}")
                        
                        if hyp.get('proposed_solution'):
                            st.markdown(f"**💡 Proposed Solution:** {hyp['proposed_solution']}")
                        
                        if hyp.get('expected_impact'):
                            st.markdown(f"**🚀 Expected Impact:** {hyp['expected_impact']}")
            else:
                st.warning("No hypotheses generated")
        
        with tab4:
            st.header("🧪 Experiment Design")
            
            # Check if we have hypotheses to work with
            if complete_data.get('hypotheses'):
                hypotheses = complete_data['hypotheses'].get('hypotheses', [])
                
                if hypotheses:
                    # Hypothesis selection
                    st.subheader("Select a Hypothesis")
                    hypothesis_options = [f"Hypothesis {h.get('hypothesis_id', i+1)}: {h.get('title', 'Untitled')}" 
                                        for i, h in enumerate(hypotheses)]
                    
                    selected_index = st.selectbox(
                        "Choose a hypothesis to generate experiments for:",
                        range(len(hypothesis_options)),
                        format_func=lambda x: hypothesis_options[x],
                        key="hypothesis_selector"
                    )
                    
                    selected_hypothesis = hypotheses[selected_index]
                    
                    # Store selected hypothesis and its gap in session state
                    st.session_state['selected_hypothesis'] = selected_hypothesis
                    st.session_state['selected_gap'] = selected_hypothesis.get('gap_addressed', 'Not specified')
                    
                    # Show selected hypothesis details
                    with st.expander("📋 Selected Hypothesis Details", expanded=False):
                        st.write(f"**Title:** {selected_hypothesis.get('title', 'N/A')}")
                        st.write(f"**Gap Addressed:** {selected_hypothesis.get('gap_addressed', 'N/A')}")
                        st.write(f"**Problem Statement:** {selected_hypothesis.get('problem_statement', 'N/A')}")
                        st.write(f"**Proposed Solution:** {selected_hypothesis.get('proposed_solution', 'N/A')}")
                        
                        col1, col2, col3 = st.columns(3)
                        with col1:
                            st.metric("Novelty Score", f"{selected_hypothesis.get('novelty_score', 'N/A')}/10")
                        with col2:
                            st.metric("Impact Score", f"{selected_hypothesis.get('impact_score', 'N/A')}/10")
                        with col3:
                            st.metric("Feasibility Score", f"{selected_hypothesis.get('feasibility_score', 'N/A')}/10")
                    
                    # Generate Experiments Button
                    if st.button("🚀 Generate Experiments & Datasets", key="generate_experiments"):
                        with st.spinner("🧪 Generating experiments, datasets, and metrics..."):
                            try:
                                experiment_results = asyncio.run(
                                    system.generate_experiments_for_hypothesis(
                                        selected_hypothesis,
                                        search_query
                                    )
                                )
                                
                                st.session_state['experiment_results'] = experiment_results
                                st.session_state['selected_experiment'] = None  # Reset selection
                                st.success("✅ Experiments generated successfully!")
                                
                            except Exception as e:
                                st.error(f"❌ Error generating experiments: {str(e)}")
                    
                    # Display experiment results
                    if st.session_state.get('experiment_results'):
                        st.markdown("---")
                        display_experiment_results(st.session_state['experiment_results'])
                else:
                    st.warning("⚠️ No hypotheses available. Please run the complete analysis first.")
            else:
                st.info("👈 Run a complete analysis first to generate hypotheses.")
        
        with tab5:
            st.header("📝 AI Research Paper Writer")
            st.markdown("*Generate a complete, publication-ready research paper in Word format*")
            
            # Check if all required data is selected
            selected_hypothesis = st.session_state.get('selected_hypothesis')
            selected_experiment = st.session_state.get('selected_experiment')
            selected_gap = st.session_state.get('selected_gap', 'Not specified')
            
            if selected_hypothesis and selected_experiment:
                # Show ready status
                st.success("✅ All required components selected! Ready to generate your paper.")
                st.markdown("---")
                
                # Show selected data summary in card-like layout
                st.markdown("### 📋 Your Research Paper Components:")
                
                col1, col2 = st.columns(2)
                
                with col1:
                    with st.container():
                        st.markdown("**📊 Research Topic**")
                        st.info(complete_data.get('topic', 'Unknown'))
                        
                        st.markdown("**💡 Hypothesis**")
                        st.info(selected_hypothesis.get('title', 'Unknown'))
                
                with col2:
                    with st.container():
                        st.markdown("**🧪 Experiment**")
                        st.info(f"{selected_experiment.get('title', 'Unknown')}")
                        
                        st.markdown("**🔍 Gap Addressed**")
                        st.info(selected_gap)
                
                st.markdown("---")
                
                # Show what will be included
                with st.expander("📄 What's included in your paper?", expanded=False):
                    col_a, col_b = st.columns(2)
                    with col_a:
                        st.markdown("""
                        **Title & Abstract**
                        - Title: 10-12 words based on your hypothesis
                        - Abstract: ~200 words summarizing the paper
                        
                        **Introduction** (~550 words, 4 paragraphs)
                        - Background on the research topic
                        - Problem statement from hypothesis
                        - Research objectives and contributions
                        - Paper structure overview
                        
                        **Literature Review** (~700 words, 4-5 paragraphs)
                        - Analysis of collected research papers
                        - Current state of the art
                        - Research gaps identified
                        - Positioning of your work
                        
                        **Methodology** (~800 words, 5-6 paragraphs)
                        - Complete experimental design
                        - Step-by-step procedures from selected experiment
                        - Datasets, metrics, and evaluation approach
                        - Implementation details
                        """)
                    with col_b:
                        st.markdown("""
                        **Expected Results** (~550 words, 4 paragraphs)
                        - Anticipated experimental outcomes
                        - Tables and figures templates
                        - Performance metrics
                        - Analysis framework
                        
                        **Discussion** (~600 words, 4-5 paragraphs)
                        - Interpretation of expected results
                        - Comparison with baselines
                        - Implications of findings
                        - Limitations and strengths
                        
                        **Conclusion** (~300 words, 3 paragraphs)
                        - Summary of contributions
                        - Key findings
                        - Future work directions
                        
                        **References**
                        - Formatted citations (APA style)
                        - All collected research papers
                        
                        **Total: ~3,500 words (4-5 pages)**
                        """)
                
                st.markdown("---")
                
                # Initialize paper writer
                writer = ResearchPaperWriter()
                
                # Context for generation
                paper_context = {
                    'topic': complete_data.get('topic', ''),
                    'hypothesis': selected_hypothesis,
                    'gap': selected_gap,
                    'experiment': selected_experiment,
                    'papers': complete_data.get('papers', []),
                    'literature_review': complete_data.get('literature_review', ''),
                    'gaps': complete_data.get('gap_analysis', {}).get('gaps', {})
                }
                
                # Generate Paper Button
                st.markdown("### 🚀 Generate Complete Research Paper")
                st.info("💡 Click the button below to generate a full research paper template. This will take 2-3 minutes.")
                
                if st.button("📝 Generate Research Paper Template", type="primary", use_container_width=True, key="generate_paper"):
                    with st.spinner("🤖 AI is writing your complete research paper... Please wait 2-3 minutes"):
                        try:
                            # Generate complete paper
                            complete_paper_md = writer.generate_complete_paper(paper_context)
                            
                            # Store in session state
                            st.session_state['generated_paper'] = complete_paper_md
                            
                            st.success("✅ Research paper generated successfully!")
                            st.rerun()
                            
                        except Exception as e:
                            st.error(f"❌ Error generating paper: {str(e)}")
                
                # Display and download generated paper
                if st.session_state.get('generated_paper'):
                    st.markdown("---")
                    st.markdown("### 📄 Your Generated Research Paper")
                    
                    # Paper stats
                    paper_text = st.session_state['generated_paper']
                    word_count = len(paper_text.split())
                    pages = word_count // 250
                    
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        st.metric("Total Words", f"{word_count:,}")
                    with col2:
                        st.metric("Est. Pages", f"~{pages}")
                    with col3:
                        st.metric("Sections", "8")
                    with col4:
                        st.metric("Status", "✅ Ready")
                    
                    st.markdown("")
                    
                    # Preview
                    with st.expander("👁️ Preview Paper Content", expanded=False):
                        st.markdown(paper_text[:3000] + "..." if len(paper_text) > 3000 else paper_text)
                    
                    st.markdown("---")
                    st.markdown("### 📥 Download Your Research Paper")
                    st.info("💡 Your paper will be downloaded as a Microsoft Word document (.docx) ready for editing and submission!")
                    
                    # Convert to Word document
                    try:
                        paper_title = selected_hypothesis.get('title', 'Research Paper')
                        doc_bytes = convert_markdown_to_docx(paper_text, paper_title)
                        
                        if doc_bytes:
                            col1, col2, col3 = st.columns([1, 2, 1])
                            with col2:
                                st.download_button(
                                    label="📄 Download as Word Document (.docx)",
                                    data=doc_bytes,
                                    file_name=f"research_paper_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    use_container_width=True,
                                    type="primary"
                                )
                        else:
                            st.error("Error converting paper to Word format. Please try again.")
                    except Exception as e:
                        st.error(f"Error creating Word document: {str(e)}")
                        st.info("💡 Tip: Make sure python-docx is installed: pip install python-docx")
                    
                    st.markdown("")
                    
                    with st.expander("ℹ️ Next Steps - How to Use Your Paper"):
                        st.markdown("""
                        ### 📝 Your Research Paper Template is Ready!
                        
                        **What's Included:**
                        - ✅ Title and Abstract
                        - ✅ Introduction with background
                        - ✅ Literature Review from collected papers
                        - ✅ Complete Methodology from your selected experiment
                        - ✅ Expected Results section (fill with actual results)
                        - ✅ Discussion framework
                        - ✅ Conclusion
                        - ✅ Formatted References
                        
                        **How to Complete Your Paper:**
                        
                        1. **Download** the Markdown or Text version
                        2. **Open** in Word, Google Docs, or any text editor
                        3. **Run** your experiment and collect results
                        4. **Add** actual results, figures, and tables to Section 4
                        5. **Update** discussion based on real findings
                        6. **Proofread** and format for your target journal
                        7. **Submit** your completed paper!
                        
                        **File Formats:**
                        - **Markdown (.md)**: Best for converting to PDF/LaTeX via Pandoc or Overleaf
                        - **Text (.txt)**: Universal format, open in any editor
                        - **JSON**: Contains paper + metadata, machine-readable
                        
                        **💡 Pro Tip:** The paper is written in academic style using GPT-4 with all your collected research context!
                        """)
                    
                    # Clear button
                    st.markdown("---")
                    if st.button("🗑️ Clear and Generate New Paper", type="secondary"):
                        st.session_state['generated_paper'] = None
                        st.success("✅ Cleared! Click 'Generate' button above to create a new paper.")
                        st.rerun()
            
            elif selected_hypothesis and not selected_experiment:
                st.warning("""
                ### ⚠️ Please Select an Experiment
                
                You have selected a hypothesis, but no experiment yet.
                
                **Steps:**
                1. Go to **Tab 4** (Experiment Design)
                2. Click **"Generate Experiments & Datasets"** if not done
                3. **Select ONE experiment** by clicking "✅ Select This Experiment"
                4. Return to this tab to generate your paper
                """)
            
            else:
                st.info("""
                ### 👈 Complete Previous Steps First
                
                To generate a research paper template, you need to:
                
                1. **Tab 1-2**: Run complete analysis to collect papers and identify gaps
                2. **Tab 3**: Review and select a hypothesis
                3. **Tab 4**: Generate experiments and **select ONE experiment**
                4. **Tab 5** (Here): Generate your complete research paper!
                
                The AI will create a professional academic paper including:
                - Title, Abstract, Introduction
                - Literature Review with citations
                - Methodology from your experiment
                - Expected Results section (for you to fill)
                - Discussion and Conclusion
                - Properly formatted References
                """)

    # Show search history
    if st.session_state.search_history:
        st.header("Search History")
        history_df = pd.DataFrame(st.session_state.search_history)
        st.dataframe(
            history_df,
            hide_index=True,
            column_config={
                "timestamp": st.column_config.DatetimeColumn(
                    "Time",
                    format="D MMM, YYYY, HH:mm"
                )
            }
        )

def display_experiment_results(results: dict):
    """Display the experiment generation results in a structured format"""
    
    st.header("🧪 Detailed Experiment Plans")
    
    # Experiments Section
    if results.get('experiments'):
        st.subheader("🔬 Select an Experiment for Your Research Paper")
        st.info("👇 Each experiment contains EVERYTHING you need: introduction, methodology, datasets, metrics, models, and challenges. Choose ONE to use in your paper.")
        
        # Experiment selection
        for i, exp in enumerate(results['experiments']):
            with st.expander(f"⚗️ Experiment {exp['id']}: {exp['title']}", expanded=(i==0)):
                
                # Header with difficulty and time
                col1, col2, col3 = st.columns([2, 1, 1])
                with col1:
                    st.markdown(f"### 📋 {exp['title']}")
                with col2:
                    difficulty_color = {"Easy": "🟢", "Medium": "🟡", "Hard": "🔴"}
                    diff = exp.get('difficulty', 'Medium')
                    st.metric("Difficulty", f"{difficulty_color.get(diff, '⚪')} {diff}")
                with col3:
                    st.metric("Estimated Time", exp.get('estimated_time', '2-3 months'))
                
                st.markdown("---")
                
                # 1. Introduction
                if exp.get('introduction'):
                    st.markdown("**📖 1. Introduction & Background**")
                    st.markdown(exp['introduction'])
                    st.markdown("")
                
                # 2. Description
                st.markdown("**🎯 2. Experiment Overview**")
                st.markdown(exp['description'])
                st.markdown("")
                
                # 3. Detailed Methodology Steps
                if exp.get('steps'):
                    st.markdown("**🔬 3. Detailed Methodology (Step-by-Step)**")
                    for j, step in enumerate(exp['steps'], 1):
                        st.markdown(f"**Step {j}:** {step}")
                    st.markdown("")
                
                # 4. Datasets (from results dict, not just experiment)
                st.markdown("**📊 4. Recommended Datasets**")
                
                if results.get('datasets'):
                    datasets = results['datasets']
                    
                    # Check if it's a list (new format) or dict (old format)
                    if isinstance(datasets, list):
                        # New format: simple list
                        if datasets:
                            for idx, dataset in enumerate(datasets[:2], 1):  # Show top 2
                                st.markdown(f"**{idx}. {dataset['name']}**")
                                st.markdown(f"   - 📦 Size: {dataset.get('size', 'N/A')}")
                                
                                # Show downloads if available
                                if dataset.get('downloads'):
                                    downloads = dataset.get('downloads')
                                    if isinstance(downloads, int):
                                        st.markdown(f"   - 💾 Downloads: {downloads:,}")
                                    else:
                                        st.markdown(f"   - 💾 Downloads: {downloads}")
                                
                                # Show usability if available
                                if dataset.get('usability'):
                                    st.markdown(f"   - ⭐ Usability: {dataset['usability']:.1f}/10")
                                
                                # Description
                                desc = dataset.get('description', '')
                                if len(desc) > 150:
                                    st.markdown(f"   - {desc[:150]}...")
                                else:
                                    st.markdown(f"   - {desc}")
                                
                                # Link
                                if dataset.get('url'):
                                    source = dataset.get('source', 'Kaggle')
                                    st.markdown(f"   - [📥 **Access Dataset**]({dataset['url']})")
                                
                                st.markdown("")
                        else:
                            st.markdown("• Custom dataset required (see methodology)")
                    else:
                        # Old format fallback: dict with kaggle/huggingface keys
                        kaggle_datasets = datasets.get('kaggle', [])
                        hf_datasets = datasets.get('huggingface', [])
                        all_datasets = kaggle_datasets + hf_datasets
                        
                        if all_datasets:
                            for idx, dataset in enumerate(all_datasets[:2], 1):
                                st.markdown(f"**{idx}. {dataset['name']}**")
                                st.markdown(f"   - {dataset.get('description', '')[:150]}")
                                if dataset.get('url'):
                                    st.markdown(f"   - [📥 **Access Dataset**]({dataset['url']})")
                                st.markdown("")
                        else:
                            st.markdown("• Custom dataset required (see methodology)")
                else:
                    st.markdown("• Dataset recommendations will be generated based on your topic")
                st.markdown("")
                
                # 5. Evaluation Metrics
                if results.get('metrics'):
                    st.markdown("**📈 5. Evaluation Metrics**")
                    for metric in results['metrics'][:4]:  # Top 4 metrics
                        st.markdown(f"• **{metric['name']}**: {metric['description']}")
                        if metric.get('range'):
                            st.markdown(f"  - Range: {metric['range']}")
                    st.markdown("")
                
                # 6. Model Architectures
                if results.get('architectures'):
                    st.markdown("**🏗️ 6. Suggested Model Architectures**")
                    for arch in results['architectures'][:3]:  # Top 3 architectures
                        st.markdown(f"• **{arch['name']}**")
                        st.markdown(f"  - {arch['description']}")
                        if arch.get('parameters'):
                            st.markdown(f"  - Parameters: {arch['parameters']}")
                    st.markdown("")
                
                # 7. Baseline Models
                if results.get('baselines'):
                    st.markdown("**📊 7. Baseline Models for Comparison**")
                    for baseline in results['baselines'][:3]:
                        st.markdown(f"• **{baseline['name']}**: {baseline['description']}")
                        if baseline.get('expected_performance'):
                            st.markdown(f"  - Expected Performance: {baseline['expected_performance']}")
                    st.markdown("")
                
                # 8. Expected Outcomes
                if results.get('expected_outcomes'):
                    st.markdown("**🎯 8. Expected Outcomes**")
                    outcomes = results['expected_outcomes']
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        st.metric("Target Accuracy", outcomes.get('accuracy_range', '85-92%'))
                    with col2:
                        st.metric("Training Time", outcomes.get('training_time', '4-6 hours'))
                    with col3:
                        st.metric("Model Size", outcomes.get('model_size', '50-100MB'))
                    with col4:
                        st.metric("Inference Speed", outcomes.get('inference_speed', '10-20ms'))
                    st.markdown("")
                
                # 9. Challenges & Solutions
                if results.get('challenges'):
                    st.markdown("**⚠️ 9. Potential Challenges & Mitigation**")
                    for challenge in results['challenges'][:4]:
                        st.markdown(f"• {challenge}")
                    st.markdown("")
                
                # 10. Required Resources
                if exp.get('required_resources'):
                    st.markdown("**💻 10. Required Resources**")
                    for resource in exp['required_resources']:
                        st.markdown(f"• {resource}")
                    st.markdown("")
                
                # Selection button - prominent
                st.markdown("---")
                col1, col2, col3 = st.columns([1, 2, 1])
                with col2:
                    if st.button(f"✓ Select This Experiment for Paper", key=f"select_exp_{exp['id']}", use_container_width=True, type="primary"):
                        st.session_state['selected_experiment'] = exp
                        st.success(f"✓ Selected: {exp['title']}")
                        st.info("→ Go to Tab 5 to generate your research paper!")
        
        # Show currently selected experiment
        if st.session_state.get('selected_experiment'):
            selected = st.session_state['selected_experiment']
            st.markdown("---")
            st.success(f"🎯 **Currently Selected:** Experiment {selected['id']} - {selected['title']}")
            st.caption("Go to Tab 5 'AI Paper Writer' to generate your complete research paper template!")
    
    # Download complete experiment plan as JSON
    st.markdown("---")
    st.subheader("💾 Export Complete Experiment Plan")
    json_str = json.dumps(results, indent=2)
    st.download_button(
        label="📥 Download All Experiments as JSON",
        data=json_str,
        file_name=f"experiment_plan_{results.get('metadata', {}).get('hypothesis_id', 'unknown')}.json",
        mime="application/json",
        help="Download all experiment details including datasets, metrics, and models"
    )

def render_footer():
    """Render footer on all pages"""
    st.markdown("<br><br>", unsafe_allow_html=True)
    st.markdown("---")
    st.markdown("""
    <div style="text-align: center; padding: 25px; background-color: #1e293b; 
                border-radius: 8px; margin-top: 30px; border-top: 2px solid #3b82f6;">
        <p style="color: #60a5fa; font-size: 13px; font-weight: 600; margin-bottom: 12px; text-transform: uppercase; letter-spacing: 1.5px;">Development Team</p>
        <p style="color: #cbd5e0; font-size: 15px; margin: 8px 0; line-height: 1.6;">
            Made by <span style="color: #f1f5f9; font-weight: 600;">Radia Riaz</span>, 
            <span style="color: #f1f5f9; font-weight: 600;">Amna Alvie</span>, 
            <span style="color: #f1f5f9; font-weight: 600;">Emaan Riaz</span>, 
            <span style="color: #f1f5f9; font-weight: 600;">Maheen Alvie</span>
        </p>
        <p style="color: #94a3b8; font-size: 12px; margin-top: 10px;">
            Gen AI Hackathon 2025
        </p>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
    render_footer()